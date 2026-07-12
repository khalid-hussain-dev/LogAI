#!/usr/bin/env python3
"""
Generate the LogAI SRS report as both HTML and PDF.

The content reflects the current implemented build while keeping performance
targets and validation evidence honest: live benchmark results can be inserted
after the final deployment test pass.
"""

from __future__ import annotations

import zipfile
from datetime import datetime
from html import escape
from pathlib import Path
from textwrap import wrap

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.shared import Cm, Pt
from PIL import Image as PILImage
from reportlab.graphics import renderPDF, renderSVG
from reportlab.graphics import renderPM
from reportlab.graphics.shapes import Circle, Drawing, Line, Polygon, Rect, String
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_LEFT
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import mm
from reportlab.platypus import (
    LongTable,
    PageBreak,
    Paragraph,
    Image as RLImage,
    SimpleDocTemplate,
    Spacer,
    Table,
    TableStyle,
)


BASE_DIR = Path(__file__).resolve().parent
HTML_PATH = BASE_DIR / "LogAI_SRS_Report.html"
PDF_PATH = BASE_DIR / "LogAI_SRS_Report.pdf"
DOCX_PATH = BASE_DIR / "LogAI_SRS_Report.docx"
DIAGRAMS_DIR = BASE_DIR / "diagrams"
TEMPLATE_PATH = BASE_DIR / "TAHAFUZ SRS REPORT (format).docx"
ASSETS_DIR = BASE_DIR / "assets"
UNIVERSITY_LOGO_PATH = ASSETS_DIR / "university_logo.png"

REPORT_DATE = datetime.now().strftime("%B %d, %Y")
PROJECT_NAME = "LogAI | AI-Powered Log Monitoring Dashboard"
REPORT_SUBTITLE = "Software Requirements Specification (SRS)"
UNIVERSITY_NAME = "Sir Syed University of Engineering & Technology"
DEPARTMENT_NAME = "Computer Engineering Department"
ACADEMIC_PROGRAM = "B.S. in Computer Engineering"
ACADEMIC_BATCH = "Batch 2022F"
REPORT_LOCATION = "Karachi, Pakistan"
INTERNAL_ADVISOR = ""
TEAM_MEMBERS = [
    ("Khalid Hussain", "2022F-BCE-130", "B", "Group Leader", "03322425697"),
    ("Nauman Khalid", "2022F-BCE-147", "A", "Member", "03295658357"),
    ("Ahmed Sartaj", "2022F-BCE-212", "A", "Member", "03303092753"),
    ("Nauroz Saleem", "2022F-BCE-150", "A", "Member", "03032809169"),
]


PALETTE = {
    "navy": colors.HexColor("#0f172a"),
    "blue": colors.HexColor("#1d4ed8"),
    "sky": colors.HexColor("#dbeafe"),
    "green": colors.HexColor("#dcfce7"),
    "mint": colors.HexColor("#bbf7d0"),
    "amber": colors.HexColor("#fef3c7"),
    "rose": colors.HexColor("#ffe4e6"),
    "gray": colors.HexColor("#475569"),
    "light": colors.HexColor("#f8fafc"),
    "border": colors.HexColor("#cbd5e1"),
    "ink": colors.HexColor("#111827"),
}


INTRODUCTION = [
    "LogAI is a web-based observability platform designed to centralize log ingestion, "
    "real-time monitoring, anomaly detection, alerting, and analytics for modern software systems.",
    "The system combines a React frontend, a FastAPI backend, a dedicated Node.js auth service, "
    "Redis-based streaming, Elasticsearch-based search and analytics, PostgreSQL-based user metadata, "
    "and containerized deployment assets.",
    "This SRS documents the current product requirements, interfaces, architecture, and design "
    "expectations for the final year project build. Performance values in this report are acceptance "
    "targets and should be validated against the live stack during the final verification run.",
]

PRODUCT_PURPOSE = [
    "The purpose of LogAI is to help operations teams, developers, and system administrators collect, "
    "observe, and analyze application logs from multiple sources in one place.",
    "The system is intended to shorten debugging time, improve operational awareness, surface anomalies "
    "quickly, and provide a cleaner workflow for log search, live monitoring, and outbound alerting.",
]

PRODUCT_SCOPE = [
    "LogAI supports secure user authentication, monitored server registration, API-key-based log ingest, "
    "Fluentd-based collection, log parsing, anomaly scoring, real-time streaming to the dashboard, search "
    "and filtering, analytics, alert integrations, and a lightweight chat assistant that explains log context.",
    "The current build also includes starter CI/CD, Docker Compose deployment assets, a reusable testing layer, "
    "and an Isolation Forest based anomaly engine with a statistical fallback path.",
]

SYSTEM_DESCRIPTION = [
    "Applications and servers submit logs either directly to the FastAPI ingest API or through Fluentd collectors.",
    "Redis Streams and Redis Lists decouple ingestion from processing so the stream worker can parse, score, and route logs reliably.",
    "The anomaly engine evaluates each log entry, then indexed logs are stored in Elasticsearch while live updates are published to Redis Pub/Sub.",
    "The React dashboard consumes historical data through REST APIs and live updates through the WebSocket endpoint.",
    "When anomalies exceed a configured threshold, outbound notifications can be dispatched through Slack, email, and generic webhooks.",
]

STAKEHOLDERS = [
    ("DevOps Engineers / SREs", "Primary users who monitor infrastructure health, incidents, and log anomalies."),
    ("Software Developers", "Investigate errors, search logs, and use analytics and chat assistance during debugging."),
    ("System Administrators", "Maintain the deployment, configure infrastructure, rotate keys, and manage environments."),
    ("Project Team", "Responsible for requirements, implementation, documentation, and future improvement of the system."),
    ("Project Advisor / Academic Evaluators", "Review system scope, architecture, correctness, testing, and deliverables."),
    ("Cloud / Hosting Providers", "Provide the runtime resources used by the databases, backend services, and frontend deployment."),
    ("Alert Delivery Services", "Slack workspaces, SMTP providers, and webhook receivers that deliver anomaly notifications."),
]

WBS_ROWS = [
    ("1.0", "Requirement analysis and proposal alignment", "Clarify scope, required features, and evaluation milestones."),
    ("2.0", "Backend platform", "Authentication, server management, ingest routes, search, analytics, and chat APIs."),
    ("3.0", "Frontend dashboard", "Connected user interface for dashboard, logs, alerts, analytics, servers, chat, and integrations."),
    ("4.0", "Real-time pipeline", "Redis streams, stream worker, WebSocket live updates, and anomaly event propagation."),
    ("5.0", "Deployment structure", "Docker Compose stack, CI/CD starter, and workspace cleanup."),
    ("6.0", "Alerting and AI/ML", "Slack/email/webhook integrations and Isolation Forest anomaly scoring."),
    ("7.0", "Testing and documentation", "Smoke tests, load tests, architecture notes, run guides, and report artifacts."),
]

GANTT_COLUMNS = ["Task", "P1", "P2", "P3", "P4", "P5", "P6", "P7"]
GANTT_ROWS = [
    ("Requirement analysis", [1, 0, 0, 0, 0, 0, 0]),
    ("Backend core implementation", [1, 1, 0, 0, 0, 0, 0]),
    ("Frontend integration and UX", [0, 1, 1, 1, 0, 0, 0]),
    ("Real-time pipeline", [0, 1, 1, 0, 0, 0, 0]),
    ("Deployment and repo structure", [0, 0, 1, 1, 0, 0, 0]),
    ("Alerting and integrations", [0, 0, 0, 0, 1, 0, 0]),
    ("AI/ML upgrade", [0, 0, 0, 0, 0, 1, 0]),
    ("Hardening, testing, documentation", [0, 0, 0, 0, 0, 0, 1]),
]

FUNCTIONAL_REQUIREMENTS = [
    ("FR-01", "User registration and login", "The system shall support local email/password signup, login, logout, token refresh, and profile retrieval.", "High"),
    ("FR-02", "OAuth callback support", "The system shall support Google and GitHub OAuth redirection through the auth service.", "Medium"),
    ("FR-03", "Server management", "The system shall allow an authenticated user to create, list, soft-delete, and rotate API keys for monitored servers.", "High"),
    ("FR-04", "Single log ingest", "The system shall accept a single log entry through an API-key-protected ingest endpoint.", "High"),
    ("FR-05", "Batch log ingest", "The system shall accept up to 1000 log entries in one batch request.", "High"),
    ("FR-06", "Collector-based ingest", "The system shall support Fluentd and Redis-based pipeline ingest in addition to direct API ingest.", "High"),
    ("FR-07", "Log parsing and normalization", "The system shall parse common raw log formats and normalize them into a shared searchable structure.", "High"),
    ("FR-08", "Real-time streaming", "The system shall push new logs and anomalies to connected dashboard clients through WebSocket updates.", "High"),
    ("FR-09", "Search and filtering", "The system shall let users query logs by server, severity, anomaly flag, text search, and time range.", "High"),
    ("FR-10", "Dashboard metrics", "The system shall present 24-hour overview metrics, severity breakdowns, hourly activity, and per-server health summaries.", "High"),
    ("FR-11", "Analytics views", "The system shall provide detailed server metrics, anomaly counts, severity charts, and trend data.", "High"),
    ("FR-12", "Alert management", "The system shall display anomaly events and support live alert updates in the frontend.", "High"),
    ("FR-13", "Outbound integrations", "The system shall allow users to configure Slack, email, and webhook channels and send test notifications.", "High"),
    ("FR-14", "AI-assisted chat", "The system shall let users ask questions about their logs and receive Elasticsearch-backed contextual responses derived from recent indexed data.", "Medium"),
    ("FR-15", "Deployment assets", "The system shall include Docker-based local deployment assets for repeatable setup and demo execution.", "Medium"),
    ("FR-16", "Testing support", "The system shall include smoke-test and ingest load-test utilities for final verification.", "Medium"),
]

OPERATING_ENVIRONMENT = [
    ("Client browser", "Chrome, Edge, or Firefox on Windows, Linux, or macOS"),
    ("Frontend runtime", "Node.js 18+ during development, static frontend served by NGINX in deployment"),
    ("Backend runtime", "Python 3.11+ for FastAPI backend and stream worker"),
    ("Auth runtime", "Node.js service for OAuth and callback handling"),
    ("Primary databases", "PostgreSQL 16, Elasticsearch 8.x, Redis 7"),
    ("Collection layer", "Fluentd with forward, HTTP, and syslog inputs"),
    ("Container environment", "Docker Desktop / Docker Compose for local stack execution"),
    ("Network prerequisites", "HTTP/HTTPS access plus database, Redis, and collector ports where applicable"),
]

UX_SCREENS = [
    ("Authentication", "Login and signup flow with local auth and OAuth callback support."),
    ("Dashboard", "Overview cards, live logs, anomaly preview, and per-server summary."),
    ("Logs", "Search, filters, pagination, and anomaly-only exploration."),
    ("Alerts", "Live anomaly view with server scoping and feed behavior."),
    ("Analytics", "Trend charts, severity breakdowns, and anomaly metrics."),
    ("Servers", "Create/list servers, view API keys, and manage monitored systems."),
    ("Integrations", "Configure Slack/email/webhook delivery and send test alerts."),
    ("Chat", "Ask log-related questions and receive contextual responses."),
]

CHAT_DETAILS = [
    "The current chat feature is not a hosted large language model integration in the deployed build.",
    "It works by querying Elasticsearch for recent logs relevant to the user's message, then composing a contextual response through rule-based response logic in the backend.",
    "Its data source is therefore the user's own indexed log history, filtered by owned servers and recent time windows.",
    "This makes the current build deterministic, lightweight, and fully local to the project stack, while still giving useful debugging assistance.",
]

PERFORMANCE_TARGETS = [
    ("Health endpoint response", "< 1 second under normal local deployment load"),
    ("Single ingest request", "< 500 ms average under light load"),
    ("Batch ingest limit", "Up to 1000 log entries per request"),
    ("Dashboard overview query", "< 3 seconds for the 24-hour aggregate view"),
    ("Log search query", "< 2 seconds for common filtered queries with paginated output"),
    ("WebSocket live update delivery", "Visible in the frontend within 2 seconds after processing"),
    ("Notification dispatch", "< 10 seconds for configured Slack/email/webhook channels"),
    ("Load-test baseline", "Support a local benchmark run of at least 200-300 ingest requests using the included test script"),
]

SAFETY_REQUIREMENTS = [
    "The system shall never perform destructive remediation actions automatically on monitored infrastructure.",
    "The system shall preserve log evidence and anomaly metadata even when live streaming is temporarily unavailable.",
    "The system shall fail gracefully by returning informative errors instead of silently dropping authenticated user actions.",
    "The system shall continue using the statistical anomaly fallback when the ML model is not trained or unavailable.",
]

SECURITY_REQUIREMENTS = [
    "JWT access and refresh tokens shall be used for authenticated user APIs.",
    "Passwords shall be stored as bcrypt hashes and never persisted in plain text.",
    "API-key-based ingest shall be isolated per monitored server.",
    "User access to servers, metrics, logs, and integrations shall be scoped to owned resources only.",
    "Rate limiting shall be applied to sensitive API flows such as ingest traffic.",
    "Secrets such as JWT keys, OAuth credentials, SMTP credentials, and webhook URLs shall be supplied through environment variables or secret stores.",
    "Production deployment should enable HTTPS/TLS at the reverse-proxy layer.",
]

BUSINESS_PERSPECTIVE = [
    "LogAI addresses a practical observability need for student projects, internal systems, and small to medium operational teams that need centralized log visibility without a large commercial stack.",
    "The product reduces manual debugging effort by combining live streams, stored history, analytics, anomaly scoring, and outbound alerting in one workflow.",
    "The platform is intentionally modular so future work can add richer ML, hosted deployment hardening, or organization-level roles without redesigning the whole system.",
]

LICENSING_REQUIREMENTS = [
    "The current build depends on open-source frameworks and libraries such as React, FastAPI, SQLAlchemy, Redis clients, Elasticsearch clients, and report/test tooling.",
    "Any public distribution of the project should include a top-level project license and third-party attribution review before release.",
    "External services such as Slack, SMTP providers, and cloud infrastructure remain subject to their own service terms and licenses.",
]

LEGAL_NOTICES = [
    "Organizations deploying LogAI remain responsible for the legality of collected log data, retention policies, and privacy obligations.",
    "The system may process sensitive operational or personal data if applications emit it into logs; therefore, deployers should implement masking, retention limits, and access controls appropriate to their context.",
    "Evidence, alerts, and exported diagnostics should be handled according to local security policies and any applicable data-protection regulations.",
]

METHODOLOGY_PARAGRAPHS = [
    "The system follows a modular, service-oriented design with clear separation between user-facing frontend flows, API orchestration, asynchronous stream processing, storage layers, and alert delivery.",
    "Development is aligned with an iterative and incremental approach: core backend services, connected frontend flows, live streaming, deployment structure, alerting, AI/ML upgrade, and final hardening were introduced in phases.",
    "The backend uses REST for historical and administrative operations, WebSocket for live updates, Redis for event movement and rate limiting, Elasticsearch for search and analytics, and PostgreSQL for user-owned metadata.",
]

DATA_DICTIONARY_RELATIONAL = [
    ("users.id", "UUID", "Primary key for each user."),
    ("users.name", "String(255)", "Display name of the user."),
    ("users.email", "String(255)", "Unique login and contact identity."),
    ("users.hashed_password", "String(255)", "Bcrypt-hashed password for local auth."),
    ("users.picture", "String(1024)", "Optional avatar URL."),
    ("users.auth_provider", "Enum", "Authentication source: local, google, or github."),
    ("users.oauth_id", "String(255)", "Optional OAuth provider identifier."),
    ("users.is_active", "Boolean", "Soft-active flag for the account."),
    ("servers.id", "UUID", "Primary key for each monitored server/application."),
    ("servers.name", "String(255)", "Display name of the monitored server."),
    ("servers.description", "Text", "Optional descriptive notes."),
    ("servers.api_key", "String(255)", "Unique API key used for log ingestion."),
    ("servers.owner_id", "UUID", "Foreign key to the owning user."),
    ("servers.is_active", "Boolean", "Soft-active flag for the server."),
    ("alert_integrations.owner_id", "UUID", "Unique foreign key linking one integration profile to one user."),
    ("alert_integrations.slack_enabled", "Boolean", "Enables Slack notifications."),
    ("alert_integrations.slack_webhook_url", "String(1024)", "Slack Incoming Webhook URL."),
    ("alert_integrations.email_enabled", "Boolean", "Enables email notifications."),
    ("alert_integrations.email_recipients", "Text", "Comma-separated email recipients."),
    ("alert_integrations.webhook_enabled", "Boolean", "Enables generic webhook delivery."),
    ("alert_integrations.webhook_url", "String(1024)", "Webhook endpoint URL."),
    ("alert_integrations.minimum_anomaly_score", "Float", "Threshold required before outbound alerting."),
]

DATA_DICTIONARY_LOG_INDEX = [
    ("logai-logs.id", "keyword", "Unique log document identifier."),
    ("logai-logs.server_id", "keyword", "Owning server identifier."),
    ("logai-logs.server_name", "keyword", "Human-readable server name."),
    ("logai-logs.level", "keyword", "Normalized severity level."),
    ("logai-logs.message", "text + keyword", "Primary searchable log message."),
    ("logai-logs.source", "keyword", "Logger or collector source."),
    ("logai-logs.host", "keyword", "Origin host name."),
    ("logai-logs.service", "keyword", "Application service or subsystem."),
    ("logai-logs.environment", "keyword", "Environment tag such as dev/staging/prod."),
    ("logai-logs.meta", "object", "Dynamic structured metadata payload."),
    ("logai-logs.raw", "text", "Original raw message when preserved."),
    ("logai-logs.anomaly", "boolean", "Boolean anomaly classification."),
    ("logai-logs.anomaly_score", "float", "0..1 anomaly score."),
    ("logai-logs.timestamp", "date", "Primary event time in epoch millis."),
    ("logai-logs.ingested_at", "date", "Time the system accepted the record."),
    ("logai-logs.processed_by", "keyword", "Pipeline source such as api-direct, stream, or fluentd."),
]

DESIGN_MODELS = [
    ("System architecture model", "Describes the high-level components and data paths from sources to dashboards and alert channels."),
    ("Use case model", "Shows how authenticated users, monitored applications, and notification channels interact with the system."),
    ("Entity relationship model", "Defines the persistent ownership relationships between users, servers, alert integrations, and log documents."),
    ("Sequence model", "Explains the order of calls during a log-ingest-to-dashboard event."),
]

EXTERNAL_INTERFACES = [
    ("POST /api/v1/auth/signup", "HTTPS", "Client -> FastAPI", "Register a local user account and receive access/refresh tokens."),
    ("POST /api/v1/auth/login", "HTTPS", "Client -> FastAPI", "Authenticate a local user and receive tokens."),
    ("POST /api/v1/auth/refresh", "HTTPS", "Client -> FastAPI", "Issue a new access token using a refresh token."),
    ("GET /api/v1/auth/me", "HTTPS", "Client -> FastAPI", "Return the authenticated user profile."),
    ("GET /api/v1/servers", "HTTPS", "Client -> FastAPI", "List owned servers with 24-hour statistics."),
    ("POST /api/v1/servers", "HTTPS", "Client -> FastAPI", "Create a monitored server and API key."),
    ("GET /api/v1/servers/{id}/metrics", "HTTPS", "Client -> FastAPI", "Fetch detailed server metrics."),
    ("GET /api/v1/servers/dashboard/overview", "HTTPS", "Client -> FastAPI", "Fetch dashboard-wide overview data."),
    ("POST /api/v1/ingest", "HTTPS + x-api-key", "Application -> FastAPI", "Submit one log entry."),
    ("POST /api/v1/ingest/batch", "HTTPS + x-api-key", "Application -> FastAPI", "Submit a batch of logs."),
    ("GET /api/v1/logs", "HTTPS", "Client -> FastAPI", "Search and filter indexed logs."),
    ("POST /api/v1/chat", "HTTPS", "Client -> FastAPI", "Ask a contextual question about current logs."),
    ("GET /api/v1/integrations/alerts", "HTTPS", "Client -> FastAPI", "Load alert integration settings."),
    ("PUT /api/v1/integrations/alerts", "HTTPS", "Client -> FastAPI", "Update alert delivery settings."),
    ("POST /api/v1/integrations/alerts/test", "HTTPS", "Client -> FastAPI", "Send a synthetic test alert."),
    ("WS /ws?token=...&server_id=...", "WebSocket", "Client <-> FastAPI", "Receive live log and anomaly events and switch subscriptions."),
    ("GET /api/auth/google, /api/auth/github", "HTTPS", "Client -> Auth Service", "Start OAuth authentication flows."),
    ("GET /api/auth/health", "HTTPS", "Client -> Auth Service", "Health endpoint for auth service validation."),
]

COMMUNICATION_INTERFACES = [
    ("Frontend to backend", "HTTPS REST + WebSocket"),
    ("Backend to PostgreSQL", "Async PostgreSQL driver"),
    ("Backend / worker to Redis", "Redis streams, lists, pub/sub, and sorted sets"),
    ("Backend / worker to Elasticsearch", "HTTP API via async client"),
    ("Fluentd to Redis", "Collector pipeline"),
    ("Alert delivery", "Slack webhook, SMTP, generic outbound webhook"),
]

ASSUMPTIONS_AND_CONSTRAINTS = [
    "The deployment environment provides Docker Desktop or an equivalent container runtime for the local stack.",
    "Node.js 18+ and Python 3.11+ are available when frontend development and local testing scripts must be run outside containers.",
    "SMTP, Slack, webhook, and OAuth behavior depends on valid third-party credentials being supplied through environment variables.",
    "The current deployment path is Docker Compose for local and demo execution.",
    "The current chat assistant is retrieval and rule based; it should not be described as a general-purpose foundation model.",
    "Final performance numbers and screenshots depend on a live system test pass and are not fully measurable from static code inspection alone.",
]

TESTING_STRATEGY = [
    ("TS-01", "Health verification", "Validate backend health and auth health endpoints before feature testing.", "Automated smoke test"),
    ("TS-02", "Authentication flow", "Verify signup, login, token refresh, logout, and OAuth callback behavior.", "Manual bug bash"),
    ("TS-03", "Server lifecycle", "Create servers, list them, and rotate keys without cross-user access leakage.", "Manual + API check"),
    ("TS-04", "Ingest validation", "Send single and batch logs with valid API keys and confirm persistence.", "Manual + smoke test"),
    ("TS-05", "WebSocket validation", "Confirm live logs and anomaly events reach the dashboard in real time.", "Manual live test"),
    ("TS-06", "Analytics and alerts", "Validate alert feed, metrics pages, and anomaly counters against ingested data.", "Manual live test"),
    ("TS-07", "Integration delivery", "Save Slack/email/webhook configuration and send a synthetic test notification.", "Manual live test"),
    ("TS-08", "Anomaly-model validation", "Warm the model with at least 100 logs and compare fallback vs ML-active behavior.", "Manual live test"),
    ("TS-09", "Load testing", "Run the ingest load test script and record throughput plus latency metrics.", "Automated load test"),
]

RISKS_AND_MITIGATIONS = [
    ("R-01", "Backend or database service failure", "Dashboard pages may fail to load or ingest may stop.", "Use health checks, container restarts, and service-level logs to detect and recover quickly."),
    ("R-02", "High ingest load", "Queue build-up, slower indexing, or delayed live updates.", "Use Redis buffering, batch indexing, load testing, and future horizontal scaling."),
    ("R-03", "Notification channel misconfiguration", "Important anomaly alerts may not be delivered.", "Provide test-send flows, readiness indicators, and per-channel configuration validation."),
    ("R-04", "Model false positives or false negatives", "Users may miss incidents or get noisy alerts.", "Keep statistical fallback, tune thresholds, and document model warm-up behavior."),
    ("R-05", "Invalid or noisy log payloads", "Parsed data quality and analytics accuracy may degrade.", "Normalize formats, support raw preservation, and store metadata for later inspection."),
    ("R-06", "Secret leakage or weak configuration", "Unauthorized access to APIs or third-party integrations.", "Use environment variables, secret stores, API key rotation, and JWT-based access control."),
]

FUTURE_IMPROVEMENTS = [
    "Upgrade the current rule-based chat assistant to a stronger retrieval-augmented AI workflow or hosted LLM integration.",
    "Explore richer anomaly models such as sequence-aware models, ensemble scoring, and feedback-driven threshold tuning.",
    "Add a hosted deployment path with production-grade monitoring, persistence, TLS, and scaling controls.",
    "Add organization-level roles, audit trails, retention policies, and multi-tenant support.",
    "Introduce formal automated E2E browser tests in addition to the current smoke and load scripts.",
]

BUG_BASH_ITEMS = [
    "Authentication and token refresh",
    "OAuth callback flow",
    "Server creation and API key rotation",
    "Single and batch ingest",
    "Live dashboard WebSocket stream",
    "Logs filtering and pagination",
    "Alerts feed behavior",
    "Analytics values and charts",
    "Integrations save and test-send flow",
    "Chat response quality for recent logs",
]


def ensure_university_logo() -> Path | None:
    if UNIVERSITY_LOGO_PATH.exists():
        return UNIVERSITY_LOGO_PATH
    if not TEMPLATE_PATH.exists():
        return None

    ASSETS_DIR.mkdir(exist_ok=True)
    with zipfile.ZipFile(TEMPLATE_PATH) as archive:
        media_files = [name for name in archive.namelist() if name.startswith("word/media/")]
        if not media_files:
            return None
        preferred = "word/media/image1.png" if "word/media/image1.png" in media_files else media_files[0]
        UNIVERSITY_LOGO_PATH.write_bytes(archive.read(preferred))
    return UNIVERSITY_LOGO_PATH


def cover_advisor_text() -> str:
    return INTERNAL_ADVISOR.strip() or "______________________________"


def cover_member_rows() -> list[list[str]]:
    return [list(member) for member in TEAM_MEMBERS]


def split_lines(text: str, width: int) -> list[str]:
    parts: list[str] = []
    for raw in text.split("\n"):
        wrapped = wrap(raw, width=width) or [""]
        parts.extend(wrapped)
    return parts


def add_box(drawing: Drawing, x: int, y: int, width: int, height: int, title: str, body: str, fill) -> None:
    drawing.add(
        Rect(
            x,
            y,
            width,
            height,
            rx=8,
            ry=8,
            fillColor=fill,
            strokeColor=PALETTE["blue"],
            strokeWidth=1.1,
        )
    )
    drawing.add(String(x + width / 2, y + height - 18, title, textAnchor="middle", fontSize=11, fillColor=PALETTE["navy"]))
    lines = split_lines(body, max(16, int(width / 7)))
    text_y = y + height - 34
    for index, line in enumerate(lines[:6]):
        drawing.add(String(x + 8, text_y - index * 12, line, fontSize=8.2, fillColor=PALETTE["ink"]))


def add_arrow(drawing: Drawing, x1: int, y1: int, x2: int, y2: int, label: str = "") -> None:
    drawing.add(Line(x1, y1, x2, y2, strokeColor=PALETTE["gray"], strokeWidth=1.2))
    angle_x = x2 - x1
    angle_y = y2 - y1
    length = (angle_x ** 2 + angle_y ** 2) ** 0.5 or 1
    ux = angle_x / length
    uy = angle_y / length
    size = 6
    left_x = x2 - ux * size - uy * size / 2
    left_y = y2 - uy * size + ux * size / 2
    right_x = x2 - ux * size + uy * size / 2
    right_y = y2 - uy * size - ux * size / 2
    drawing.add(
        Polygon(
            [x2, y2, left_x, left_y, right_x, right_y],
            fillColor=PALETTE["gray"],
            strokeColor=PALETTE["gray"],
        )
    )
    if label:
        drawing.add(String((x1 + x2) / 2, (y1 + y2) / 2 + 6, label, fontSize=7.5, fillColor=PALETTE["gray"], textAnchor="middle"))


def system_diagram() -> Drawing:
    d = Drawing(860, 470)
    add_box(d, 25, 355, 160, 70, "Applications / Servers", "Custom apps, shippers,\nand monitored services", PALETTE["sky"])
    add_box(d, 25, 245, 160, 70, "Fluentd Collectors", "Forward, HTTP, and\nsyslog input paths", PALETTE["amber"])
    add_box(d, 220, 355, 170, 70, "FastAPI Ingest API", "Single and batch ingest\nprotected by API keys", PALETTE["green"])
    add_box(d, 220, 245, 170, 70, "Redis Stream / List", "Durable event buffering\nfor asynchronous processing", PALETTE["amber"])
    add_box(d, 220, 120, 170, 80, "Auth Service + PostgreSQL", "JWT/OAuth flows,\nusers, servers, and alert settings", PALETTE["rose"])
    add_box(d, 430, 280, 170, 90, "Stream Worker + Parser", "Normalize logs, parse\nformats, and coordinate scoring", PALETTE["sky"])
    add_box(d, 430, 155, 170, 90, "Anomaly + Notification Engine", "Isolation Forest with\nstatistical fallback and outbound alerts", PALETTE["green"])
    add_box(d, 640, 310, 180, 80, "Elasticsearch", "Search, analytics,\nand dashboard history", PALETTE["sky"])
    add_box(d, 640, 205, 180, 70, "Redis Pub/Sub + WebSocket", "Live logs and anomalies\nfor connected dashboards", PALETTE["amber"])
    add_box(d, 640, 95, 180, 70, "React Frontend", "Dashboard, logs, alerts,\nanalytics, chat, integrations", PALETTE["green"])
    add_box(d, 640, 10, 180, 55, "Slack / Email / Webhook", "Outbound anomaly alerts", PALETTE["rose"])

    add_arrow(d, 185, 390, 220, 390, "direct ingest")
    add_arrow(d, 185, 280, 220, 280, "collector stream")
    add_arrow(d, 305, 355, 305, 315, "")
    add_arrow(d, 390, 280, 430, 325, "consume")
    add_arrow(d, 390, 390, 430, 325, "score path")
    add_arrow(d, 515, 280, 515, 245, "")
    add_arrow(d, 600, 325, 640, 350, "index")
    add_arrow(d, 600, 235, 640, 235, "publish")
    add_arrow(d, 730, 205, 730, 165, "stream")
    add_arrow(d, 600, 195, 640, 40, "notify")
    add_arrow(d, 390, 155, 430, 195, "auth data")
    add_arrow(d, 390, 160, 640, 130, "REST")
    return d


def screen_map_diagram() -> Drawing:
    d = Drawing(860, 300)
    d.add(Rect(30, 30, 800, 230, rx=10, ry=10, fillColor=PALETTE["light"], strokeColor=PALETTE["border"]))
    d.add(Rect(30, 30, 165, 230, fillColor=PALETTE["navy"], strokeColor=PALETTE["navy"]))
    for index, label in enumerate(["Dashboard", "Logs", "Alerts", "Analytics", "Servers", "Chat", "Integrations", "Docs", "Settings"]):
        d.add(String(52, 228 - index * 22, label, fontSize=9.5, fillColor=colors.white))
    d.add(Rect(215, 200, 590, 35, fillColor=colors.white, strokeColor=PALETTE["border"]))
    d.add(String(235, 213, "Overview Header + Filters", fontSize=10, fillColor=PALETTE["gray"]))
    for x in [225, 420, 615]:
        d.add(Rect(x, 145, 160, 40, fillColor=PALETTE["sky"], strokeColor=PALETTE["blue"]))
    d.add(String(255, 162, "KPI Card", fontSize=9))
    d.add(String(450, 162, "KPI Card", fontSize=9))
    d.add(String(645, 162, "KPI Card", fontSize=9))
    d.add(Rect(225, 65, 360, 65, fillColor=PALETTE["green"], strokeColor=PALETTE["blue"]))
    d.add(String(365, 100, "Charts / Trends / Live Activity", fontSize=10, textAnchor="middle"))
    d.add(Rect(605, 65, 180, 65, fillColor=PALETTE["amber"], strokeColor=PALETTE["blue"]))
    d.add(String(695, 100, "Live Logs / Alerts Panel", fontSize=10, textAnchor="middle"))
    return d


def use_case_diagram() -> Drawing:
    d = Drawing(860, 370)
    d.add(Circle(70, 260, 22, strokeColor=PALETTE["navy"], fillColor=None))
    d.add(Line(70, 238, 70, 190, strokeColor=PALETTE["navy"]))
    d.add(Line(50, 225, 90, 225, strokeColor=PALETTE["navy"]))
    d.add(Line(70, 190, 50, 160, strokeColor=PALETTE["navy"]))
    d.add(Line(70, 190, 90, 160, strokeColor=PALETTE["navy"]))
    d.add(String(70, 145, "User", textAnchor="middle", fontSize=10))

    d.add(Circle(70, 80, 22, strokeColor=PALETTE["navy"], fillColor=None))
    d.add(Line(70, 58, 70, 20, strokeColor=PALETTE["navy"]))
    d.add(Line(50, 45, 90, 45, strokeColor=PALETTE["navy"]))
    d.add(Line(70, 20, 50, -10, strokeColor=PALETTE["navy"]))
    d.add(Line(70, 20, 90, -10, strokeColor=PALETTE["navy"]))
    d.add(String(70, -25, "External App", textAnchor="middle", fontSize=10))

    add_box(d, 190, 255, 165, 52, "Authenticate", "Sign up, login,\nrefresh tokens", PALETTE["sky"])
    add_box(d, 385, 255, 165, 52, "Manage Servers", "Create servers,\nrotate keys", PALETTE["green"])
    add_box(d, 580, 255, 165, 52, "View Dashboard", "Read KPIs, live logs,\nand server health", PALETTE["amber"])
    add_box(d, 190, 165, 165, 52, "Search Logs", "Filter, paginate,\nand inspect events", PALETTE["green"])
    add_box(d, 385, 165, 165, 52, "Inspect Alerts", "Review anomalies\nand active events", PALETTE["rose"])
    add_box(d, 580, 165, 165, 52, "Analytics + Chat", "Charts, trends,\nand log Q&A", PALETTE["sky"])
    add_box(d, 385, 75, 165, 52, "Configure Integrations", "Slack, email,\nwebhook settings", PALETTE["amber"])
    add_box(d, 190, 20, 165, 52, "Send Logs", "Single or batch\nAPI ingest", PALETTE["green"])
    add_box(d, 580, 20, 165, 52, "Receive Notifications", "Outbound anomaly\ndelivery", PALETTE["rose"])

    for target_x, target_y in [(190, 280), (385, 280), (580, 280), (190, 190), (385, 190), (580, 190), (385, 100)]:
        add_arrow(d, 95, 225, target_x, target_y, "")
    add_arrow(d, 95, 80, 190, 45, "")
    add_arrow(d, 745, 45, 820, 45, "")
    d.add(String(810, 60, "Alert Channel", textAnchor="middle", fontSize=10))
    return d


def er_diagram() -> Drawing:
    d = Drawing(860, 340)
    add_box(d, 40, 200, 220, 110, "users", "id (UUID)\nname\nemail\nhashed_password\nauth_provider\ncreated_at / updated_at", PALETTE["sky"])
    add_box(d, 320, 200, 220, 110, "servers", "id (UUID)\nname\ndescription\napi_key\nowner_id -> users.id\nis_active", PALETTE["green"])
    add_box(d, 600, 200, 220, 110, "alert_integrations", "owner_id -> users.id\nslack/email/webhook flags\nrecipient URLs\nminimum_anomaly_score", PALETTE["rose"])
    add_box(d, 320, 40, 220, 110, "logai-logs (Elasticsearch)", "id\nserver_id\nlevel\nmessage\nanomaly\nanomaly_score\ntimestamp", PALETTE["amber"])
    add_arrow(d, 260, 255, 320, 255, "1 : N")
    add_arrow(d, 150, 200, 150, 160, "")
    add_arrow(d, 150, 160, 600, 255, "1 : 1")
    add_arrow(d, 430, 200, 430, 150, "1 : N")
    d.add(String(150, 145, "one user owns one alert profile", fontSize=8, textAnchor="middle"))
    d.add(String(430, 135, "one server produces many log documents", fontSize=8, textAnchor="middle"))
    return d


def sequence_diagram() -> Drawing:
    d = Drawing(860, 390)
    actors = [
        ("Application", 60),
        ("Ingest API", 180),
        ("Anomaly Engine", 320),
        ("Elasticsearch", 470),
        ("Redis", 610),
        ("WebSocket", 730),
        ("Frontend", 820),
    ]
    for label, x in actors:
        d.add(String(x, 360, label, textAnchor="middle", fontSize=10, fillColor=PALETTE["navy"]))
        d.add(Line(x, 40, x, 340, strokeColor=PALETTE["border"], strokeDashArray=[4, 3]))

    steps = [
        (60, 180, 320, "1. POST /ingest"),
        (180, 320, 290, "2. score(log)"),
        (320, 470, 260, "3. store indexed log"),
        (320, 610, 220, "4. publish live event"),
        (610, 730, 180, "5. broadcast pub/sub"),
        (730, 820, 140, "6. push WebSocket message"),
        (320, 610, 100, "7. optional alert dispatch"),
    ]
    for x1, x2, y, label in steps:
        add_arrow(d, x1, y, x2, y, label)
    return d


def drawing_svg(drawing: Drawing) -> str:
    svg = renderSVG.drawToString(drawing)
    if isinstance(svg, bytes):
        return svg.decode("utf-8")
    return str(svg)


def export_diagram_assets() -> None:
    DIAGRAMS_DIR.mkdir(parents=True, exist_ok=True)
    diagrams = {
        "system_diagram": system_diagram(),
        "ux_prototype": screen_map_diagram(),
        "use_case_diagram": use_case_diagram(),
        "entity_relationship_diagram": er_diagram(),
        "sequence_diagram": sequence_diagram(),
    }
    for name, drawing in diagrams.items():
        svg_path = DIAGRAMS_DIR / f"{name}.svg"
        png_path = DIAGRAMS_DIR / f"{name}.png"
        renderSVG.drawToFile(drawing, str(svg_path))
        renderPM.drawToFile(drawing, str(png_path), fmt="PNG")


def diagram_png(name: str) -> Path:
    return DIAGRAMS_DIR / f"{name}.png"


def rl_diagram_image(name: str, max_width_mm: float = 170, max_height_mm: float = 190):
    path = diagram_png(name)
    with PILImage.open(path) as img:
        width_px, height_px = img.size
    max_width_pt = max_width_mm * mm
    max_height_pt = max_height_mm * mm
    scale = min(max_width_pt / width_px, max_height_pt / height_px, 1.0)
    return RLImage(str(path), width=width_px * scale, height=height_px * scale)


def html_table(headers: list[str], rows: list[list[str]]) -> str:
    parts = ["<table><thead><tr>"]
    parts.extend(f"<th>{escape(header)}</th>" for header in headers)
    parts.append("</tr></thead><tbody>")
    for row in rows:
        parts.append("<tr>")
        parts.extend(f"<td>{escape(str(cell))}</td>" for cell in row)
        parts.append("</tr>")
    parts.append("</tbody></table>")
    return "".join(parts)


def html_gantt_table() -> str:
    parts = ["<table class='gantt'><thead><tr>"]
    for header in GANTT_COLUMNS:
        parts.append(f"<th>{escape(header)}</th>")
    parts.append("</tr></thead><tbody>")
    for label, states in GANTT_ROWS:
        parts.append(f"<tr><td>{escape(label)}</td>")
        for state in states:
            cell_class = "active" if state else "idle"
            cell_text = "Active" if state else ""
            parts.append(f"<td class='{cell_class}'>{cell_text}</td>")
        parts.append("</tr>")
    parts.append("</tbody></table>")
    return "".join(parts)


def html_list(items: list[str]) -> str:
    return "<ul>" + "".join(f"<li>{escape(item)}</li>" for item in items) + "</ul>"


def html_stakeholders() -> str:
    return html_table(["Stakeholder", "Role / Interest"], [list(item) for item in STAKEHOLDERS])


def html_functional_requirements() -> str:
    return html_table(["ID", "Requirement", "Description", "Priority"], [list(item) for item in FUNCTIONAL_REQUIREMENTS])


def html_operating_environment() -> str:
    return html_table(["Environment Item", "Current Build Expectation"], [list(item) for item in OPERATING_ENVIRONMENT])


def html_performance_targets() -> str:
    return html_table(["Metric", "Acceptance Target"], [list(item) for item in PERFORMANCE_TARGETS])


def html_data_dictionary() -> str:
    relational = html_table(["Field", "Type", "Description"], [list(item) for item in DATA_DICTIONARY_RELATIONAL])
    logs = html_table(["Field", "Type", "Description"], [list(item) for item in DATA_DICTIONARY_LOG_INDEX])
    return (
        "<h4>Relational entities</h4>"
        + relational
        + "<h4>Log index dataset</h4>"
        + logs
    )


def html_design_models() -> str:
    return html_table(["Model", "Purpose"], [list(item) for item in DESIGN_MODELS])


def html_external_interfaces() -> str:
    api_table = html_table(["Interface", "Protocol", "Direction", "Description"], [list(item) for item in EXTERNAL_INTERFACES])
    comms = html_table(["Communication Path", "Technology"], [list(item) for item in COMMUNICATION_INTERFACES])
    return "<h4>API and service interfaces</h4>" + api_table + "<h4>Communication interfaces</h4>" + comms


def html_wbs_table() -> str:
    return html_table(["WBS", "Work Package", "Description"], [list(item) for item in WBS_ROWS])


def html_screen_cards() -> str:
    cards = []
    for title, desc in UX_SCREENS:
        cards.append(
            "<div class='screen-card'>"
            f"<h4>{escape(title)}</h4>"
            f"<p>{escape(desc)}</p>"
            "</div>"
        )
    return "<div class='screen-grid'>" + "".join(cards) + "</div>"


def build_html() -> str:
    system_svg = drawing_svg(system_diagram())
    ux_svg = drawing_svg(screen_map_diagram())
    use_case_svg = drawing_svg(use_case_diagram())
    er_svg = drawing_svg(er_diagram())
    sequence_svg = drawing_svg(sequence_diagram())
    logo_path = ensure_university_logo()
    logo_src = f"assets/{logo_path.name}" if logo_path else ""

    intro_html = "".join(f"<p>{escape(p)}</p>" for p in INTRODUCTION)
    purpose_html = "".join(f"<p>{escape(p)}</p>" for p in PRODUCT_PURPOSE)
    scope_html = "".join(f"<p>{escape(p)}</p>" for p in PRODUCT_SCOPE)
    system_desc_html = html_list(SYSTEM_DESCRIPTION)
    methodology_html = "".join(f"<p>{escape(p)}</p>" for p in METHODOLOGY_PARAGRAPHS)
    business_html = html_list(BUSINESS_PERSPECTIVE)
    safety_html = html_list(SAFETY_REQUIREMENTS)
    security_html = html_list(SECURITY_REQUIREMENTS)
    licensing_html = html_list(LICENSING_REQUIREMENTS)
    legal_html = html_list(LEGAL_NOTICES)
    bug_bash_html = html_list(BUG_BASH_ITEMS)
    assumptions_html = html_list(ASSUMPTIONS_AND_CONSTRAINTS)
    chat_html = html_list(CHAT_DETAILS)
    testing_html = html_table(["ID", "Validation Area", "Expected Check", "Method"], [list(item) for item in TESTING_STRATEGY])
    risks_html = html_table(["Risk ID", "Risk", "Impact", "Mitigation"], [list(item) for item in RISKS_AND_MITIGATIONS])
    future_html = html_list(FUTURE_IMPROVEMENTS)
    team_html = html_table(["Name", "Roll No", "Section", "Role", "Contact"], cover_member_rows())

    return f"""<!doctype html>
<html lang="en">
<head>
  <meta charset="utf-8" />
  <meta name="viewport" content="width=device-width, initial-scale=1" />
  <title>{escape(PROJECT_NAME)} - SRS Report</title>
  <style>
    :root {{
      --navy: #0f172a;
      --blue: #1d4ed8;
      --sky: #dbeafe;
      --green: #dcfce7;
      --amber: #fef3c7;
      --rose: #ffe4e6;
      --light: #f8fafc;
      --border: #cbd5e1;
      --ink: #111827;
      --muted: #475569;
    }}
    * {{ box-sizing: border-box; }}
    body {{
      margin: 0;
      background: linear-gradient(180deg, #eef4ff 0%, #f8fafc 32%, #ffffff 100%);
      color: var(--ink);
      font-family: "Segoe UI", Tahoma, Geneva, Verdana, sans-serif;
      line-height: 1.6;
    }}
    .page {{
      max-width: 1120px;
      margin: 0 auto;
      padding: 42px 32px 80px;
    }}
    .hero {{
      background: radial-gradient(circle at top left, #dbeafe 0%, #ffffff 55%);
      border: 1px solid var(--border);
      border-radius: 24px;
      padding: 36px;
      box-shadow: 0 18px 45px rgba(15, 23, 42, 0.08);
      margin-bottom: 28px;
    }}
    .hero h1 {{
      margin: 0 0 6px;
      font-size: 2.4rem;
      color: var(--navy);
    }}
    .hero h2 {{
      margin: 0;
      font-size: 1.2rem;
      font-weight: 600;
      color: var(--blue);
    }}
    .hero-logo {{
      width: 110px;
      height: 110px;
      object-fit: contain;
      display: block;
      margin: 0 auto 18px;
    }}
    .hero-kicker {{
      text-transform: uppercase;
      letter-spacing: 0.08em;
      color: var(--muted);
      font-size: 0.92rem;
      margin-bottom: 8px;
      text-align: center;
    }}
    .hero-meta {{
      max-width: 720px;
      margin: 18px auto 0;
      text-align: center;
      color: var(--navy);
    }}
    .hero-meta p {{
      margin: 4px 0;
    }}
    .hero-members {{
      margin-top: 20px;
    }}
    .hero-members table {{
      margin-bottom: 0;
      font-size: 0.92rem;
    }}
    .hero-members th {{
      background: #e7eefb;
    }}
    .hero p {{
      margin: 12px 0 0;
      color: var(--muted);
    }}
    .toc {{
      background: rgba(255, 255, 255, 0.9);
      border: 1px solid var(--border);
      border-radius: 20px;
      padding: 22px 24px;
      margin-bottom: 26px;
    }}
    .toc h3, section h3 {{
      margin-top: 0;
      color: var(--navy);
    }}
    .toc a {{
      color: var(--blue);
      text-decoration: none;
    }}
    .toc ol {{ margin: 0; padding-left: 20px; }}
    section {{
      background: rgba(255, 255, 255, 0.95);
      border: 1px solid var(--border);
      border-radius: 20px;
      padding: 26px;
      margin-bottom: 20px;
      box-shadow: 0 10px 30px rgba(15, 23, 42, 0.05);
    }}
    h3 {{
      font-size: 1.5rem;
      margin-bottom: 12px;
    }}
    h4 {{
      margin: 22px 0 8px;
      color: var(--blue);
    }}
    p {{ margin: 0 0 10px; }}
    ul {{ margin-top: 8px; }}
    .note {{
      border-left: 5px solid var(--blue);
      background: #eef5ff;
      padding: 14px 16px;
      border-radius: 12px;
      color: var(--navy);
      margin: 14px 0 20px;
    }}
    .figure {{
      background: var(--light);
      border: 1px solid var(--border);
      border-radius: 16px;
      padding: 18px;
      margin: 16px 0 10px;
      overflow: auto;
    }}
    .caption {{
      font-size: 0.92rem;
      color: var(--muted);
      margin-bottom: 18px;
    }}
    table {{
      width: 100%;
      border-collapse: collapse;
      margin: 12px 0 20px;
      font-size: 0.95rem;
    }}
    th, td {{
      border: 1px solid var(--border);
      padding: 10px 12px;
      vertical-align: top;
    }}
    th {{
      background: #eff6ff;
      color: var(--navy);
      text-align: left;
    }}
    .screen-grid {{
      display: grid;
      grid-template-columns: repeat(auto-fit, minmax(220px, 1fr));
      gap: 14px;
      margin-top: 16px;
    }}
    .screen-card {{
      background: linear-gradient(180deg, #ffffff 0%, #f8fafc 100%);
      border: 1px solid var(--border);
      border-radius: 16px;
      padding: 16px;
    }}
    .screen-card h4 {{ margin-top: 0; }}
    .gantt td.active {{
      background: #bfdbfe;
      color: #0f172a;
      font-weight: 600;
      text-align: center;
    }}
    .gantt td.idle {{
      background: #ffffff;
    }}
    @media print {{
      body {{ background: white; }}
      .page {{ padding: 0; max-width: none; }}
      .hero, .toc, section {{
        box-shadow: none;
        border-color: #d6dbe3;
      }}
      a {{ color: inherit; }}
    }}
  </style>
</head>
<body>
  <div class="page">
    <div class="hero">
      {f'<img class="hero-logo" src="{escape(logo_src)}" alt="University Logo" />' if logo_src else ''}
      <div class="hero-kicker">{escape(UNIVERSITY_NAME)}</div>
      <h1>{escape(REPORT_SUBTITLE)}</h1>
      <h2>{escape(PROJECT_NAME)}</h2>
      <div class="hero-meta">
        <p><strong>{escape(DEPARTMENT_NAME)}</strong></p>
        <p>{escape(ACADEMIC_PROGRAM)}, {escape(ACADEMIC_BATCH)}</p>
        <p><strong>Internal Advisor:</strong> {escape(cover_advisor_text())}</p>
        <p><strong>Date:</strong> {escape(REPORT_DATE)}</p>
      </div>
      <div class="hero-members">
        <p><strong>Project Team</strong></p>
        {team_html}
      </div>
    </div>

    <div class="toc">
      <h3>Contents</h3>
      <ol>
        <li><a href="#intro">1.1 Introduction</a></li>
        <li><a href="#purpose">1.1.1 Product Purpose</a></li>
        <li><a href="#scope">1.1.2 Product Scope</a></li>
        <li><a href="#system">1.1.3 System Diagram and Description</a></li>
        <li><a href="#stakeholders">1.1.4 Project Stakeholders</a></li>
        <li><a href="#wbs">1.1.5 Project WBS and Gantt Chart</a></li>
        <li><a href="#assumptions">1.1.6 Assumptions and Constraints</a></li>
        <li><a href="#functional">1.2 Project Functional Requirements</a></li>
        <li><a href="#functional">1.2.1.1 Chat Assistant Details</a></li>
        <li><a href="#nonfunctional">1.3 Project Non-Functional Requirements</a></li>
        <li><a href="#nonfunctional">1.3.7 Verification and Testing Strategy</a></li>
        <li><a href="#nonfunctional">1.3.8 Risks and Mitigation</a></li>
        <li><a href="#methodology">1.4 Software Design Methodology</a></li>
        <li><a href="#interfaces">1.5 External Interface Requirements</a></li>
        <li><a href="#er">1.6 Entity Relationship Diagram</a></li>
        <li><a href="#sequence">1.7 Sequence Diagram</a></li>
        <li><a href="#future">1.8 Future Improvements</a></li>
      </ol>
    </div>

    <section id="intro">
      <h3>1.1 Introduction</h3>
      {intro_html}
      <h4 id="purpose">1.1.1 Product Purpose</h4>
      {purpose_html}
      <h4 id="scope">1.1.2 Product Scope</h4>
      {scope_html}
      <div class="note">
        This report documents the implemented product requirements and the expected final verification criteria.
        Sections such as performance evidence and final screenshots should be updated after the live test run.
      </div>
    </section>

    <section id="system">
      <h3>1.1.3 System Diagram and Description</h3>
      <div class="figure">{system_svg}</div>
      <div class="caption">Figure 1: LogAI system diagram</div>
      {system_desc_html}
    </section>

    <section id="stakeholders">
      <h3>1.1.4 Project Stakeholders</h3>
      {html_stakeholders()}
    </section>

    <section id="wbs">
      <h3>1.1.5 Project WBS and Gantt Chart</h3>
      <h4>Work Breakdown Structure</h4>
      {html_wbs_table()}
      <h4>Phase-wise Gantt overview</h4>
      {html_gantt_table()}
      <div class="caption">P1 to P7 correspond to the seven project phases used to structure the implementation roadmap.</div>
    </section>

    <section id="assumptions">
      <h3>1.1.6 Assumptions and Constraints</h3>
      {assumptions_html}
    </section>

    <section id="functional">
      <h3>1.2 Project Functional Requirements</h3>
      <h4>1.2.1 Project Functional Features</h4>
      {html_functional_requirements()}
      <h4>1.2.1.1 Chat Assistant Details</h4>
      {chat_html}
      <h4>1.2.2 Operating Environment</h4>
      {html_operating_environment()}
      <h4>1.2.3 UX Prototype</h4>
      <div class="figure">{ux_svg}</div>
      <div class="caption">Figure 2: Representative frontend layout prototype based on the current page structure</div>
      {html_screen_cards()}
    </section>

    <section id="nonfunctional">
      <h3>1.3 Project Non-Functional Requirements</h3>
      <h4>1.3.1 Performance Evaluation</h4>
      <div class="note">The following values are acceptance targets for the final verification pass. Actual results should be inserted after running the live smoke and load tests.</div>
      {html_performance_targets()}
      <h4>1.3.2 Safety Requirements</h4>
      {safety_html}
      <h4>1.3.3 Security Requirements</h4>
      {security_html}
      <h4>1.3.4 Business Perspective / Requirements</h4>
      {business_html}
      <h4>1.3.5 Licensing Requirements</h4>
      {licensing_html}
      <h4>1.3.6 Legal Copyrights and Notices</h4>
      {legal_html}
      <h4>1.3.7 Verification and Testing Strategy</h4>
      {testing_html}
      <h4>1.3.8 Risks and Mitigation</h4>
      {risks_html}
    </section>

    <section id="methodology">
      <h3>1.4 Software Design Methodology</h3>
      {methodology_html}
      <h4>1.4.1 Data Dictionary / Data Set</h4>
      {html_data_dictionary()}
      <h4>1.4.2 Design Models [along with descriptions]</h4>
      {html_design_models()}
      <div class="figure">{use_case_svg}</div>
      <div class="caption">Figure 3: Use case diagram for the current build</div>
    </section>

    <section id="interfaces">
      <h3>1.5 External Interface Requirements</h3>
      {html_external_interfaces()}
      <h4>Recommended manual bug-bash coverage</h4>
      {bug_bash_html}
    </section>

    <section id="er">
      <h3>1.6 Entity Relationship Diagram</h3>
      <div class="figure">{er_svg}</div>
      <div class="caption">Figure 4: ER diagram combining relational ownership and the Elasticsearch log dataset</div>
    </section>

    <section id="sequence">
      <h3>1.7 Sequence Diagram</h3>
      <div class="figure">{sequence_svg}</div>
      <div class="caption">Figure 5: Sequence diagram for a log ingest to live dashboard event</div>
    </section>

    <section id="future">
      <h3>1.8 Future Improvements</h3>
      {future_html}
    </section>
  </div>
</body>
</html>"""


def styles():
    base = getSampleStyleSheet()
    base.add(
        ParagraphStyle(
            name="ReportTitle",
            fontName="Times-Bold",
            fontSize=20,
            leading=25,
            alignment=TA_CENTER,
            textColor=colors.black,
            spaceAfter=10,
        )
    )
    base.add(
        ParagraphStyle(
            name="ReportSubtitle",
            fontName="Times-Bold",
            fontSize=28,
            leading=34,
            alignment=TA_CENTER,
            textColor=colors.black,
            spaceAfter=10,
        )
    )
    base.add(
        ParagraphStyle(
            name="SectionHeading",
            fontName="Times-Bold",
            fontSize=14,
            leading=18,
            textColor=colors.black,
            spaceBefore=14,
            spaceAfter=8,
        )
    )
    base.add(
        ParagraphStyle(
            name="SubHeading",
            fontName="Times-Bold",
            fontSize=11.5,
            leading=14,
            textColor=colors.black,
            spaceBefore=10,
            spaceAfter=6,
        )
    )
    base.add(
        ParagraphStyle(
            name="BodyTextSmall",
            parent=base["BodyText"],
            fontName="Times-Roman",
            fontSize=11,
            leading=15,
            textColor=colors.black,
            alignment=TA_LEFT,
        )
    )
    base.add(
        ParagraphStyle(
            name="Note",
            parent=base["BodyText"],
            fontName="Times-Italic",
            fontSize=10,
            leading=13.5,
            backColor=colors.HexColor("#f7f7f7"),
            borderColor=colors.black,
            borderPadding=6,
            borderWidth=0.6,
            borderRadius=6,
            spaceBefore=6,
            spaceAfter=10,
        )
    )
    return base


def pdf_table(headers: list[str], rows: list[list[str]], col_widths=None, font_size: float = 8.4):
    data = [headers] + rows
    table = LongTable(data, colWidths=col_widths, repeatRows=1)
    table.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#f0f0f0")),
                ("TEXTCOLOR", (0, 0), (-1, 0), colors.black),
                ("FONTNAME", (0, 0), (-1, 0), "Times-Bold"),
                ("FONTNAME", (0, 1), (-1, -1), "Times-Roman"),
                ("FONTSIZE", (0, 0), (-1, -1), font_size),
                ("LEADING", (0, 0), (-1, -1), font_size + 2),
                ("GRID", (0, 0), (-1, -1), 0.5, colors.black),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#fafafa")]),
                ("LEFTPADDING", (0, 0), (-1, -1), 5),
                ("RIGHTPADDING", (0, 0), (-1, -1), 5),
                ("TOPPADDING", (0, 0), (-1, -1), 4),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
            ]
        )
    )
    return table


def pdf_gantt_table():
    headers = GANTT_COLUMNS
    rows: list[list[str]] = []
    for label, states in GANTT_ROWS:
        rows.append([label] + [("Active" if item else "") for item in states])
    widths = [75 * mm] + [14 * mm] * 7
    table = Table([headers] + rows, colWidths=widths, repeatRows=1)
    style_commands = [
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#f0f0f0")),
        ("FONTNAME", (0, 0), (-1, 0), "Times-Bold"),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.black),
        ("GRID", (0, 0), (-1, -1), 0.5, colors.black),
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
        ("FONTSIZE", (0, 0), (-1, -1), 7.8),
        ("ALIGN", (1, 1), (-1, -1), "CENTER"),
    ]
    for row_index, (_, states) in enumerate(GANTT_ROWS, start=1):
        for col_index, state in enumerate(states, start=1):
            if state:
                style_commands.append(("BACKGROUND", (col_index, row_index), (col_index, row_index), colors.HexColor("#d9d9d9")))
                style_commands.append(("TEXTCOLOR", (col_index, row_index), (col_index, row_index), colors.black))
                style_commands.append(("FONTNAME", (col_index, row_index), (col_index, row_index), "Times-Bold"))
    table.setStyle(TableStyle(style_commands))
    return table


def bullet_paragraphs(items: list[str], style) -> list[Paragraph]:
    return [Paragraph(f"- {escape(item)}", style) for item in items]


def add_paragraphs(story: list, paragraphs: list[str], style) -> None:
    for paragraph in paragraphs:
        story.append(Paragraph(escape(paragraph), style))
        story.append(Spacer(1, 4))


def add_header_footer(canvas, doc) -> None:
    canvas.saveState()
    canvas.setFont("Helvetica", 8)
    canvas.setFillColor(PALETTE["gray"])
    canvas.drawString(doc.leftMargin, 12 * mm, PROJECT_NAME)
    canvas.drawRightString(doc.pagesize[0] - doc.rightMargin, 12 * mm, f"Page {doc.page}")
    canvas.restoreState()


def build_pdf() -> None:
    doc = SimpleDocTemplate(
        str(PDF_PATH),
        pagesize=A4,
        rightMargin=18 * mm,
        leftMargin=18 * mm,
        topMargin=16 * mm,
        bottomMargin=18 * mm,
    )
    s = styles()
    story: list = []

    contents_items = [
        "1.1 Introduction",
        "1.1.1 Product Purpose",
        "1.1.2 Product Scope",
        "1.1.3 System Diagram and Description",
        "1.1.4 Project Stakeholders",
        "1.1.5 Project WBS and Gantt Chart",
        "1.1.6 Assumptions and Constraints",
        "1.2 Project Functional Requirements",
        "1.2.1.1 Chat Assistant Details",
        "1.3 Project Non-Functional Requirements",
        "1.3.7 Verification and Testing Strategy",
        "1.3.8 Risks and Mitigation",
        "1.4 Software Design Methodology",
        "1.5 External Interface Requirements",
        "1.6 Entity Relationship Diagram",
        "1.7 Sequence Diagram",
        "1.8 Future Improvements",
    ]
    figure_items = [
        "Figure 1.1: LogAI System Diagram",
        "Figure 1.2: UI/UX Prototype Overview",
        "Figure 1.3: Use Case Diagram",
        "Figure 1.4: Entity Relationship Diagram",
        "Figure 1.5: Sequence Diagram",
    ]

    def add_center_line(text: str, style_name: str, spacer_after: float = 3):
        story.append(Paragraph(text, s[style_name]))
        story.append(Spacer(1, spacer_after))

    logo_path = ensure_university_logo()
    story.append(Spacer(1, 24))
    if logo_path:
        cover_logo = RLImage(str(logo_path), width=30 * mm, height=30 * mm)
        cover_logo.hAlign = "CENTER"
        story.append(cover_logo)
        story.append(Spacer(1, 8))
    add_center_line(UNIVERSITY_NAME.upper(), "BodyTextSmall", 2)
    add_center_line(DEPARTMENT_NAME.upper(), "BodyTextSmall", 12)
    add_center_line("Software Requirement Specification", "ReportTitle", 6)
    add_center_line(PROJECT_NAME, "ReportSubtitle", 6)
    add_center_line(f"{ACADEMIC_PROGRAM}, {ACADEMIC_BATCH}", "BodyTextSmall", 10)
    add_center_line(f"<b>Internal Advisor:</b> {escape(cover_advisor_text())}", "BodyTextSmall", 14)
    add_center_line("<b>Project Team</b>", "BodyTextSmall", 6)
    team_table = pdf_table(
        ["Name", "Roll No", "Section", "Role", "Contact"],
        cover_member_rows(),
        col_widths=[45 * mm, 35 * mm, 18 * mm, 34 * mm, 42 * mm],
        font_size=9.1,
    )
    team_table.hAlign = "CENTER"
    story.append(team_table)
    story.append(Spacer(1, 10))
    add_center_line(REPORT_LOCATION, "BodyTextSmall", 2)
    add_center_line(REPORT_DATE, "BodyTextSmall", 0)
    story.append(PageBreak())

    story.append(Paragraph("Contents", s["SectionHeading"]))
    for item in contents_items:
        story.append(Paragraph(item, s["BodyTextSmall"]))
    story.append(Spacer(1, 10))
    story.append(Paragraph("List of Figures", s["SectionHeading"]))
    for item in figure_items:
        story.append(Paragraph(item, s["BodyTextSmall"]))
    story.append(PageBreak())

    story.append(Paragraph("1.1 Introduction", s["SectionHeading"]))
    add_paragraphs(story, INTRODUCTION, s["BodyTextSmall"])
    story.append(Paragraph("1.1.1 Product Purpose", s["SubHeading"]))
    add_paragraphs(story, PRODUCT_PURPOSE, s["BodyTextSmall"])
    story.append(Paragraph("1.1.2 Product Scope", s["SubHeading"]))
    add_paragraphs(story, PRODUCT_SCOPE, s["BodyTextSmall"])
    story.append(Paragraph("Performance numbers and final screenshots should be updated after the live deployment test pass.", s["Note"]))
    story.append(PageBreak())

    story.append(Paragraph("1.1.3 System Diagram and Description", s["SectionHeading"]))
    img = rl_diagram_image("system_diagram", max_width_mm=168, max_height_mm=155)
    img.hAlign = "CENTER"
    story.append(img)
    story.append(Spacer(1, 5))
    story.append(Paragraph("Figure 1.1: LogAI System Diagram", s["BodyTextSmall"]))
    story.extend(bullet_paragraphs(SYSTEM_DESCRIPTION, s["BodyTextSmall"]))
    story.append(Spacer(1, 8))
    story.append(Paragraph("1.1.4 Project Stakeholders", s["SectionHeading"]))
    story.append(pdf_table(["Stakeholder", "Role / Interest"], [list(item) for item in STAKEHOLDERS], col_widths=[58 * mm, 116 * mm], font_size=9.2))
    story.append(Spacer(1, 6))
    story.append(Paragraph("1.1.5 Project WBS and Gantt Chart", s["SectionHeading"]))
    story.append(Paragraph("Work Breakdown Structure", s["SubHeading"]))
    story.append(pdf_table(["WBS", "Work Package", "Description"], [list(item) for item in WBS_ROWS], col_widths=[18 * mm, 58 * mm, 98 * mm], font_size=8.8))
    story.append(Spacer(1, 6))
    story.append(Paragraph("Phase-wise Gantt overview", s["SubHeading"]))
    story.append(pdf_gantt_table())
    story.append(Spacer(1, 8))
    story.append(Paragraph("1.1.6 Assumptions and Constraints", s["SectionHeading"]))
    story.extend(bullet_paragraphs(ASSUMPTIONS_AND_CONSTRAINTS, s["BodyTextSmall"]))
    story.append(PageBreak())

    story.append(Paragraph("1.2 Project Functional Requirements", s["SectionHeading"]))
    story.append(Paragraph("1.2.1 Project Functional Features", s["SubHeading"]))
    story.append(pdf_table(["ID", "Requirement", "Description", "Priority"], [list(item) for item in FUNCTIONAL_REQUIREMENTS], col_widths=[18 * mm, 43 * mm, 105 * mm, 18 * mm], font_size=8.1))
    story.append(Spacer(1, 6))
    story.append(Paragraph("1.2.1.1 Chat Assistant Details", s["SubHeading"]))
    story.extend(bullet_paragraphs(CHAT_DETAILS, s["BodyTextSmall"]))
    story.append(Spacer(1, 6))
    story.append(Paragraph("1.2.2 Operating Environment", s["SubHeading"]))
    story.append(pdf_table(["Environment Item", "Current Build Expectation"], [list(item) for item in OPERATING_ENVIRONMENT], col_widths=[58 * mm, 116 * mm], font_size=9.0))
    story.append(Spacer(1, 6))
    story.append(Paragraph("1.2.3 UX Prototype", s["SubHeading"]))
    ux_img = rl_diagram_image("ux_prototype", max_width_mm=165, max_height_mm=120)
    ux_img.hAlign = "CENTER"
    story.append(ux_img)
    story.append(Paragraph("Figure 1.2: UI/UX Prototype Overview", s["BodyTextSmall"]))
    for title, desc in UX_SCREENS:
        story.append(Paragraph(f"<b>{escape(title)}</b>: {escape(desc)}", s["BodyTextSmall"]))
    story.append(PageBreak())

    story.append(Paragraph("1.3 Project Non-Functional Requirements", s["SectionHeading"]))
    story.append(Paragraph("1.3.1 Performance Evaluation", s["SubHeading"]))
    story.append(Paragraph("The following values are acceptance targets pending the final live verification run.", s["Note"]))
    story.append(pdf_table(["Metric", "Acceptance Target"], [list(item) for item in PERFORMANCE_TARGETS], col_widths=[70 * mm, 104 * mm], font_size=9.0))
    story.append(Paragraph("1.3.2 Safety Requirements", s["SubHeading"]))
    story.extend(bullet_paragraphs(SAFETY_REQUIREMENTS, s["BodyTextSmall"]))
    story.append(Spacer(1, 6))
    story.append(Paragraph("1.3.3 Security Requirements", s["SubHeading"]))
    story.extend(bullet_paragraphs(SECURITY_REQUIREMENTS, s["BodyTextSmall"]))
    story.append(Spacer(1, 6))
    story.append(Paragraph("1.3.4 Business Perspective / Requirements", s["SubHeading"]))
    story.extend(bullet_paragraphs(BUSINESS_PERSPECTIVE, s["BodyTextSmall"]))
    story.append(Spacer(1, 6))
    story.append(Paragraph("1.3.5 Licensing Requirements", s["SubHeading"]))
    story.extend(bullet_paragraphs(LICENSING_REQUIREMENTS, s["BodyTextSmall"]))
    story.append(Spacer(1, 6))
    story.append(Paragraph("1.3.6 Legal Copyrights and Notices", s["SubHeading"]))
    story.extend(bullet_paragraphs(LEGAL_NOTICES, s["BodyTextSmall"]))
    story.append(Spacer(1, 6))
    story.append(Paragraph("1.3.7 Verification and Testing Strategy", s["SubHeading"]))
    story.append(pdf_table(["ID", "Validation Area", "Expected Check", "Method"], [list(item) for item in TESTING_STRATEGY], col_widths=[16 * mm, 40 * mm, 88 * mm, 30 * mm], font_size=8.0))
    story.append(Spacer(1, 6))
    story.append(Paragraph("1.3.8 Risks and Mitigation", s["SubHeading"]))
    story.append(pdf_table(["Risk ID", "Risk", "Impact", "Mitigation"], [list(item) for item in RISKS_AND_MITIGATIONS], col_widths=[16 * mm, 44 * mm, 44 * mm, 70 * mm], font_size=8.0))
    story.append(PageBreak())

    story.append(Paragraph("1.4 Software Design Methodology", s["SectionHeading"]))
    add_paragraphs(story, METHODOLOGY_PARAGRAPHS, s["BodyTextSmall"])
    story.append(Paragraph("1.4.1 Data Dictionary / Data Set", s["SubHeading"]))
    story.append(Paragraph("Relational entities", s["BodyTextSmall"]))
    story.append(pdf_table(["Field", "Type", "Description"], [list(item) for item in DATA_DICTIONARY_RELATIONAL], col_widths=[48 * mm, 33 * mm, 93 * mm], font_size=8.0))
    story.append(Spacer(1, 6))
    story.append(Paragraph("Log index dataset", s["BodyTextSmall"]))
    story.append(pdf_table(["Field", "Type", "Description"], [list(item) for item in DATA_DICTIONARY_LOG_INDEX], col_widths=[48 * mm, 28 * mm, 98 * mm], font_size=8.0))
    story.append(Spacer(1, 6))
    story.append(Paragraph("1.4.2 Design Models [along with descriptions]", s["SubHeading"]))
    story.append(pdf_table(["Model", "Purpose"], [list(item) for item in DESIGN_MODELS], col_widths=[58 * mm, 116 * mm], font_size=9.0))
    story.append(PageBreak())

    story.append(Paragraph("1.4.2.1 Use Case Diagram", s["SubHeading"]))
    uc_img = rl_diagram_image("use_case_diagram", max_width_mm=168, max_height_mm=160)
    uc_img.hAlign = "CENTER"
    story.append(uc_img)
    story.append(Paragraph("Figure 1.3: Use Case Diagram", s["BodyTextSmall"]))
    story.append(PageBreak())

    story.append(Paragraph("1.5 External Interface Requirements", s["SectionHeading"]))
    story.append(Paragraph("API and service interfaces", s["SubHeading"]))
    story.append(pdf_table(["Interface", "Protocol", "Direction", "Description"], [list(item) for item in EXTERNAL_INTERFACES], col_widths=[46 * mm, 22 * mm, 32 * mm, 74 * mm], font_size=7.8))
    story.append(Spacer(1, 6))
    story.append(Paragraph("Communication interfaces", s["SubHeading"]))
    story.append(pdf_table(["Communication Path", "Technology"], [list(item) for item in COMMUNICATION_INTERFACES], col_widths=[68 * mm, 106 * mm], font_size=9.0))
    story.append(Spacer(1, 6))
    story.append(Paragraph("Recommended manual bug-bash coverage", s["SubHeading"]))
    story.extend(bullet_paragraphs(BUG_BASH_ITEMS, s["BodyTextSmall"]))
    story.append(PageBreak())

    story.append(Paragraph("1.6 Entity Relationship Diagram", s["SectionHeading"]))
    er_img = rl_diagram_image("entity_relationship_diagram", max_width_mm=165, max_height_mm=155)
    er_img.hAlign = "CENTER"
    story.append(er_img)
    story.append(Paragraph("Figure 1.4: Entity Relationship Diagram", s["BodyTextSmall"]))
    story.append(PageBreak())

    story.append(Paragraph("1.7 Sequence Diagram", s["SectionHeading"]))
    seq_img = rl_diagram_image("sequence_diagram", max_width_mm=168, max_height_mm=165)
    seq_img.hAlign = "CENTER"
    story.append(seq_img)
    story.append(Paragraph("Figure 1.5: Sequence Diagram", s["BodyTextSmall"]))
    story.append(Spacer(1, 8))
    story.append(Paragraph("1.8 Future Improvements", s["SectionHeading"]))
    story.extend(bullet_paragraphs(FUTURE_IMPROVEMENTS, s["BodyTextSmall"]))

    doc.build(story, onFirstPage=add_header_footer, onLaterPages=add_header_footer)


def clear_document_body(doc: Document) -> None:
    body = doc._element.body
    for child in list(body):
        if child.tag.endswith("sectPr"):
            continue
        body.remove(child)


def set_run_font(run, size_pt: int | float | None = None, bold: bool | None = None) -> None:
    run.font.name = "Times New Roman"
    if size_pt is not None:
        run.font.size = Pt(size_pt)
    if bold is not None:
        run.bold = bold


def add_docx_paragraph(doc: Document, text: str, style: str = "Normal", align=None, size_pt=None, bold=None, after_space=6):
    para = doc.add_paragraph(style=style)
    if align is not None:
        para.alignment = align
    run = para.add_run(text)
    set_run_font(run, size_pt=size_pt, bold=bold)
    para.paragraph_format.space_after = Pt(after_space)
    return para


def add_docx_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        para = doc.add_paragraph(style="List Paragraph")
        para.paragraph_format.space_after = Pt(4)
        run = para.add_run(item)
        set_run_font(run, size_pt=11)


def add_docx_table(doc: Document, headers: list[str], rows: list[list[str]]):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr_cells = table.rows[0].cells
    for idx, header in enumerate(headers):
        hdr_cells[idx].text = header
        for paragraph in hdr_cells[idx].paragraphs:
            for run in paragraph.runs:
                set_run_font(run, size_pt=10.5, bold=True)
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = str(value)
            for paragraph in cells[idx].paragraphs:
                for run in paragraph.runs:
                    set_run_font(run, size_pt=10.5)
    doc.add_paragraph("")
    return table


def add_docx_image(doc: Document, image_name: str, width_cm: float, caption: str) -> None:
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run()
    run.add_picture(str(diagram_png(image_name)), width=Cm(width_cm))
    cap = doc.add_paragraph(style="Body Text")
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap_run = cap.add_run(caption)
    set_run_font(cap_run, size_pt=10.5)


def build_docx() -> None:
    doc = Document(str(TEMPLATE_PATH)) if TEMPLATE_PATH.exists() else Document()
    clear_document_body(doc)

    section = doc.sections[0]
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)

    normal_style = doc.styles["Normal"]
    normal_style.font.name = "Times New Roman"
    normal_style.font.size = Pt(12)
    try:
        doc.styles["Body Text"].font.name = "Times New Roman"
        doc.styles["Body Text"].font.size = Pt(12)
    except Exception:
        pass

    logo_path = ensure_university_logo()
    if logo_path:
        logo_para = doc.add_paragraph()
        logo_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        logo_run = logo_para.add_run()
        logo_run.add_picture(str(logo_path), width=Cm(3.4))

    add_docx_paragraph(doc, UNIVERSITY_NAME.upper(), align=WD_ALIGN_PARAGRAPH.CENTER, size_pt=16, bold=True, after_space=2)
    add_docx_paragraph(doc, DEPARTMENT_NAME.upper(), align=WD_ALIGN_PARAGRAPH.CENTER, size_pt=14, bold=True, after_space=14)
    add_docx_paragraph(doc, "Software Requirement Specification", align=WD_ALIGN_PARAGRAPH.CENTER, size_pt=18, bold=True, after_space=4)
    add_docx_paragraph(doc, PROJECT_NAME, align=WD_ALIGN_PARAGRAPH.CENTER, size_pt=30, bold=True, after_space=8)
    add_docx_paragraph(doc, f"{ACADEMIC_PROGRAM}, {ACADEMIC_BATCH}", align=WD_ALIGN_PARAGRAPH.CENTER, size_pt=14, bold=True, after_space=12)
    add_docx_paragraph(doc, f"Internal Advisor: {cover_advisor_text()}", align=WD_ALIGN_PARAGRAPH.CENTER, size_pt=13, bold=True, after_space=12)
    add_docx_paragraph(doc, "Project Team", align=WD_ALIGN_PARAGRAPH.CENTER, size_pt=14, bold=True, after_space=4)
    team_table = add_docx_table(doc, ["Name", "Roll No", "Section", "Role", "Contact"], cover_member_rows())
    team_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    add_docx_paragraph(doc, REPORT_LOCATION, align=WD_ALIGN_PARAGRAPH.CENTER, size_pt=12, after_space=2)
    add_docx_paragraph(doc, REPORT_DATE, align=WD_ALIGN_PARAGRAPH.CENTER, size_pt=12, after_space=0)
    doc.add_page_break()

    add_docx_paragraph(doc, "Contents", style="Heading 1", size_pt=16, bold=True)
    for item in [
        "1.1 Introduction",
        "1.1.1 Product Purpose",
        "1.1.2 Product Scope",
        "1.1.3 System Diagram and Description",
        "1.1.4 Project Stakeholders",
        "1.1.5 Project WBS and Gantt Chart",
        "1.1.6 Assumptions and Constraints",
        "1.2 Project Functional Requirements",
        "1.2.1.1 Chat Assistant Details",
        "1.3 Project Non-Functional Requirements",
        "1.3.7 Verification and Testing Strategy",
        "1.3.8 Risks and Mitigation",
        "1.4 Software Design Methodology",
        "1.5 External Interface Requirements",
        "1.6 Entity Relationship Diagram",
        "1.7 Sequence Diagram",
        "1.8 Future Improvements",
    ]:
        add_docx_paragraph(doc, item, style="Body Text", size_pt=11, after_space=2)
    add_docx_paragraph(doc, "List of Figures", style="Heading 1", size_pt=16, bold=True)
    for item in [
        "Figure 1.1: LogAI System Diagram",
        "Figure 1.2: UI/UX Prototype Overview",
        "Figure 1.3: Use Case Diagram",
        "Figure 1.4: Entity Relationship Diagram",
        "Figure 1.5: Sequence Diagram",
    ]:
        add_docx_paragraph(doc, item, style="Body Text", size_pt=11, after_space=2)

    add_docx_paragraph(doc, "1.1 Introduction", style="Heading 2", size_pt=14, bold=True)
    for paragraph in INTRODUCTION:
        add_docx_paragraph(doc, paragraph, style="Body Text", size_pt=11)
    add_docx_paragraph(doc, "1.1.1 Product Purpose", style="Heading 3", size_pt=12, bold=True)
    for paragraph in PRODUCT_PURPOSE:
        add_docx_paragraph(doc, paragraph, style="Body Text", size_pt=11)
    add_docx_paragraph(doc, "1.1.2 Product Scope", style="Heading 3", size_pt=12, bold=True)
    for paragraph in PRODUCT_SCOPE:
        add_docx_paragraph(doc, paragraph, style="Body Text", size_pt=11)
    add_docx_paragraph(doc, "1.1.3 System Diagram and Description", style="Heading 3", size_pt=12, bold=True)
    add_docx_image(doc, "system_diagram", 16.5, "Figure 1.1: LogAI System Diagram")
    add_docx_bullets(doc, SYSTEM_DESCRIPTION)
    add_docx_paragraph(doc, "1.1.4 Project Stakeholders", style="Heading 3", size_pt=12, bold=True)
    add_docx_table(doc, ["Stakeholder", "Role / Interest"], [list(item) for item in STAKEHOLDERS])
    add_docx_paragraph(doc, "1.1.5 Project WBS and Gantt Chart", style="Heading 3", size_pt=12, bold=True)
    add_docx_table(doc, ["WBS", "Work Package", "Description"], [list(item) for item in WBS_ROWS])
    add_docx_table(doc, GANTT_COLUMNS, [[label] + [("Active" if s else "") for s in states] for label, states in GANTT_ROWS])
    add_docx_paragraph(doc, "1.1.6 Assumptions and Constraints", style="Heading 3", size_pt=12, bold=True)
    add_docx_bullets(doc, ASSUMPTIONS_AND_CONSTRAINTS)

    doc.add_page_break()
    add_docx_paragraph(doc, "1.2 Project Functional Requirements", style="Heading 2", size_pt=14, bold=True)
    add_docx_paragraph(doc, "1.2.1 Project Functional Features", style="Heading 3", size_pt=12, bold=True)
    add_docx_table(doc, ["ID", "Requirement", "Description", "Priority"], [list(item) for item in FUNCTIONAL_REQUIREMENTS])
    add_docx_paragraph(doc, "1.2.1.1 Chat Assistant Details", style="Heading 3", size_pt=12, bold=True)
    add_docx_bullets(doc, CHAT_DETAILS)
    add_docx_paragraph(doc, "1.2.2 Operating Environment", style="Heading 3", size_pt=12, bold=True)
    add_docx_table(doc, ["Environment Item", "Current Build Expectation"], [list(item) for item in OPERATING_ENVIRONMENT])
    add_docx_paragraph(doc, "1.2.3 UX Prototype", style="Heading 3", size_pt=12, bold=True)
    add_docx_image(doc, "ux_prototype", 16.0, "Figure 1.2: UI/UX Prototype Overview")
    for title, desc in UX_SCREENS:
        add_docx_paragraph(doc, f"{title}: {desc}", style="Body Text", size_pt=11)

    doc.add_page_break()
    add_docx_paragraph(doc, "1.3 Project Non-Functional Requirements", style="Heading 2", size_pt=14, bold=True)
    add_docx_paragraph(doc, "1.3.1 Performance Evaluation", style="Heading 3", size_pt=12, bold=True)
    add_docx_paragraph(doc, "The following values are acceptance targets pending the final live verification run.", style="Body Text", size_pt=11)
    add_docx_table(doc, ["Metric", "Acceptance Target"], [list(item) for item in PERFORMANCE_TARGETS])
    add_docx_paragraph(doc, "1.3.2 Safety Requirements", style="Heading 3", size_pt=12, bold=True)
    add_docx_bullets(doc, SAFETY_REQUIREMENTS)
    add_docx_paragraph(doc, "1.3.3 Security Requirements", style="Heading 3", size_pt=12, bold=True)
    add_docx_bullets(doc, SECURITY_REQUIREMENTS)
    add_docx_paragraph(doc, "1.3.4 Business Perspective / Requirements", style="Heading 3", size_pt=12, bold=True)
    add_docx_bullets(doc, BUSINESS_PERSPECTIVE)
    add_docx_paragraph(doc, "1.3.5 Licensing Requirements", style="Heading 3", size_pt=12, bold=True)
    add_docx_bullets(doc, LICENSING_REQUIREMENTS)
    add_docx_paragraph(doc, "1.3.6 Legal Copyrights and Notices", style="Heading 3", size_pt=12, bold=True)
    add_docx_bullets(doc, LEGAL_NOTICES)
    add_docx_paragraph(doc, "1.3.7 Verification and Testing Strategy", style="Heading 3", size_pt=12, bold=True)
    add_docx_table(doc, ["ID", "Validation Area", "Expected Check", "Method"], [list(item) for item in TESTING_STRATEGY])
    add_docx_paragraph(doc, "1.3.8 Risks and Mitigation", style="Heading 3", size_pt=12, bold=True)
    add_docx_table(doc, ["Risk ID", "Risk", "Impact", "Mitigation"], [list(item) for item in RISKS_AND_MITIGATIONS])

    doc.add_page_break()
    add_docx_paragraph(doc, "1.4 Software Design Methodology", style="Heading 2", size_pt=14, bold=True)
    for paragraph in METHODOLOGY_PARAGRAPHS:
        add_docx_paragraph(doc, paragraph, style="Body Text", size_pt=11)
    add_docx_paragraph(doc, "1.4.1 Data Dictionary / Data Set", style="Heading 3", size_pt=12, bold=True)
    add_docx_table(doc, ["Field", "Type", "Description"], [list(item) for item in DATA_DICTIONARY_RELATIONAL])
    add_docx_table(doc, ["Field", "Type", "Description"], [list(item) for item in DATA_DICTIONARY_LOG_INDEX])
    add_docx_paragraph(doc, "1.4.2 Design Models [along with descriptions]", style="Heading 3", size_pt=12, bold=True)
    add_docx_table(doc, ["Model", "Purpose"], [list(item) for item in DESIGN_MODELS])
    add_docx_image(doc, "use_case_diagram", 16.2, "Figure 1.3: Use Case Diagram")

    doc.add_page_break()
    add_docx_paragraph(doc, "1.5 External Interface Requirements", style="Heading 2", size_pt=14, bold=True)
    add_docx_table(doc, ["Interface", "Protocol", "Direction", "Description"], [list(item) for item in EXTERNAL_INTERFACES])
    add_docx_table(doc, ["Communication Path", "Technology"], [list(item) for item in COMMUNICATION_INTERFACES])
    add_docx_paragraph(doc, "Recommended manual bug-bash coverage", style="Heading 3", size_pt=12, bold=True)
    add_docx_bullets(doc, BUG_BASH_ITEMS)
    add_docx_paragraph(doc, "1.6 Entity Relationship Diagram", style="Heading 2", size_pt=14, bold=True)
    add_docx_image(doc, "entity_relationship_diagram", 16.0, "Figure 1.4: Entity Relationship Diagram")
    add_docx_paragraph(doc, "1.7 Sequence Diagram", style="Heading 2", size_pt=14, bold=True)
    add_docx_image(doc, "sequence_diagram", 16.2, "Figure 1.5: Sequence Diagram")
    add_docx_paragraph(doc, "1.8 Future Improvements", style="Heading 2", size_pt=14, bold=True)
    add_docx_bullets(doc, FUTURE_IMPROVEMENTS)

    doc.save(str(DOCX_PATH))


def main() -> None:
    export_diagram_assets()
    HTML_PATH.write_text(build_html(), encoding="utf-8")
    build_docx()
    build_pdf()
    print(f"Generated HTML: {HTML_PATH}")
    print(f"Generated DOCX: {DOCX_PATH}")
    print(f"Generated PDF:  {PDF_PATH}")
    print(f"Generated diagrams in: {DIAGRAMS_DIR}")


if __name__ == "__main__":
    main()
